"""
DOCX Parser
Extracts text from Word document attachments with comprehensive error handling
"""

import io
import logging
import os
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)


class DOCXParser:
    """
    Parser for extracting text from DOCX attachments
    Uses python-docx library for Word document processing
    """

    def __init__(self):
        self.supported_mimetypes = [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/docx",
        ]

    def extract_text(self, docx_path: str) -> str:
        """
        Extract text from a DOCX file
        Returns the complete text content
        """
        from docx import Document

        text = ""

        if not os.path.exists(docx_path):
            logger.error(f"DOCX file not found: {docx_path}")
            return ""

        file_size = os.path.getsize(docx_path)
        if file_size > 50 * 1024 * 1024:  # 50MB limit
            logger.warning(f"DOCX file too large: {file_size / (1024*1024):.1f}MB")
            return ""

        try:
            doc = Document(docx_path)

            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"

            extracted_length = len(text)
            logger.info(f"Extracted {extracted_length} characters from DOCX: {docx_path}")
            return text.strip()

        except Exception as e:
            logger.error(f"Error parsing DOCX {docx_path}: {e}")
            return ""

    def extract_text_from_bytes(self, docx_bytes: bytes, max_size_mb: int = 50) -> str:
        """
        Extract text from DOCX bytes (when file is not saved)
        Returns the complete text content
        
        Args:
            docx_bytes: Raw DOCX bytes
            max_size_mb: Maximum file size in MB
        """
        from docx import Document

        text = ""

        if not docx_bytes:
            logger.warning("Empty DOCX bytes provided")
            return ""

        size_mb = len(docx_bytes) / (1024 * 1024)
        if size_mb > max_size_mb:
            logger.warning(f"DOCX size ({size_mb:.1f}MB) exceeds limit ({max_size_mb}MB)")
            return ""

        try:
            doc = Document(io.BytesIO(docx_bytes))

            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"

            logger.info(f"Extracted {len(text)} characters from DOCX bytes ({size_mb:.2f}MB)")
            return text.strip()

        except Exception as e:
            logger.error(f"Error parsing DOCX from bytes: {e}")
            return ""

    def extract_tables(self, docx_path: str) -> List[List[List[str]]]:
        """
        Extract tables from DOCX (if any exist)
        Returns list of tables, each table is a list of rows
        """
        tables = []

        try:
            doc = Document(docx_path)

            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                tables.append(table_data)

            return tables

        except Exception as e:
            print(f"❌ Error extracting tables from DOCX: {e}")
            return []

    def get_docx_info(self, docx_path: str) -> Dict:
        """
        Get metadata information about the DOCX
        """
        info = {}

        try:
            doc = Document(docx_path)

            # Core properties
            if doc.core_properties.author:
                info["author"] = doc.core_properties.author
            if doc.core_properties.created:
                info["created"] = doc.core_properties.created.isoformat()
            if doc.core_properties.modified:
                info["modified"] = doc.core_properties.modified.isoformat()
            if doc.core_properties.title:
                info["title"] = doc.core_properties.title
            if doc.core_properties.subject:
                info["subject"] = doc.core_properties.subject

            return info

        except Exception as e:
            print(f"❌ Error getting DOCX info: {e}")
            return info

    def is_possible_order_docx(self, text: str) -> bool:
        """
        Check if DOCX text appears to contain order-related content
        """
        order_keywords = [
            "purchase order",
            "po #",
            "po no",
            "order no",
            "order date",
            "delivery date",
            "quantity",
            "product",
            "item",
            "supplier",
            "customer",
            "invoice",
            "proforma",
            "enquiry",
            "inquiry",
        ]

        text_lower = text.lower()
        keyword_count = sum(1 for keyword in order_keywords if keyword in text_lower)

        # If 3 or more order keywords found, likely an order document
        return keyword_count >= 3

    def parse_docx_attachment(self, attachment: Dict, attachment_path: str) -> Dict:
        """
        Parse a DOCX attachment and return extracted data
        """
        result = {
            "filename": attachment.get("filename", ""),
            "mimetype": attachment.get("mimeType", ""),
            "text": "",
            "tables": [],
            "is_order_docx": False,
            "success": False,
            "error": None,
        }

        try:
            # Extract text
            text = self.extract_text(attachment_path)
            result["text"] = text
            result["is_order_docx"] = self.is_possible_order_docx(text)
            result["success"] = True

        except Exception as e:
            result["error"] = str(e)

        return result


# Test function
def test_docx_parser():
    """Test the DOCX parser"""
    parser = DOCXParser()

    print("🧪 Testing DOCX Parser\n" + "=" * 50)

    # Test with a sample DOCX (if exists)
    test_docx_path = "tests/sample_order.docx"

    if os.path.exists(test_docx_path):
        print(f"\n📄 Parsing: {test_docx_path}")

        # Extract text
        text = parser.extract_text(test_docx_path)
        print(f"\n📊 Extracted text (first 200 chars):")
        print(text[:200] + "..." if len(text) > 200 else text)

        # Get DOCX info
        info = parser.get_docx_info(test_docx_path)
        print(f"\n📝 DOCX Info:")
        for key, value in info.items():
            print(f"   {key}: {value}")

        # Check if it's an order DOCX
        is_order = parser.is_possible_order_docx(text)
        print(f"\n✅ Order DOCX: {is_order}")

    else:
        print(f"\nℹ️  No test DOCX found at {test_docx_path}")
        print("   Create a sample DOCX to test parsing.")

    print("\n" + "=" * 50)
    print("✅ DOCX Parser testing complete!")


if __name__ == "__main__":
    test_docx_parser()
